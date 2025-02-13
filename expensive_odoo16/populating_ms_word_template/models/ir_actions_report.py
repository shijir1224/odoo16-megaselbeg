# -*- coding: utf-8 -*-

from odoo import models, fields, tools
from odoo.addons.populating_ms_word_template.models.mailmerge import MailMerge
import binascii
import tempfile
from zipfile import ZipFile, ZIP_DEFLATED
from io import BytesIO
from datetime import datetime,date
import logging
import pytz
import sys
import subprocess
import re
from lxml import etree

from docx import Document
from docx.shared import Inches
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from odoo.tools.image import base64_to_image
import base64

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'ct': 'http://schemas.openxmlformats.org/package/2006/content-types',
}


_logger = logging.getLogger(__name__)
try:
    from num2words import num2words
except ImportError:
    _logger.warning("The num2words python library is not installed, amount-to-text features won't be fully available.")
    num2words = None


class IrActionsReport(models.Model):
    _inherit = 'ir.actions.report'

    file_template_data = fields.Binary('File template', attachment=True)
    file_template_name = fields.Char('File template')
    populating_ms_word_template = fields.Boolean('Populating MSWord Template', default=False)
    type_export = fields.Selection([('pdf','PDF'), ('docx','Docx')], string='Export', default='docx')

    def _convert_binary_to_doc(self, file_template_data=None, suffix='.docx'):
        fp = tempfile.NamedTemporaryFile(suffix=suffix)
        if file_template_data == None:
            fp.write(binascii.a2b_base64(self.file_template_data))
        else:
            fp.write(binascii.a2b_base64(file_template_data))
        fp.seek(0)
        return fp

    def export_doc_by_template(self, file_template_data=None, suffix='.docx', file_name_export='export1', datas={}):
        simple_merge = {}
        populating_tables = {}
        file_template = self._convert_binary_to_doc(file_template_data=file_template_data,suffix='.docx')
        document = MailMerge(file_template.name)
        fields = document.get_merge_fields()

        for field in fields:
            childs = field.split('.')
            if len(childs) == 1:
                value = getattr(datas, childs[0], '')
                if isinstance(value, datetime):
                    value = self._convert_datetime_usertz_to_utctz(value)
                elif isinstance(value, date):
                    value = value.strftime(self.env['res.lang'].search([('code', '=', self.env.user.lang)], limit=1).date_format)
                elif isinstance(value, bool):
                    if value == False:
                        value = ''
                    else:
                        value = str(value)
                elif isinstance(value, bytes):
                    value = childs[0]
                else:
                    if isinstance(value, int) or isinstance(value, float):
                        value = self.format_number(value)
                    else:
                        value = str(value)
                simple_merge[field] = value
            else:
                if childs[0] == 'line' and len(childs)>1 and childs[1] == 'user_sign':
                    childs.remove(childs[0])
                    childs.remove(childs[0])
                    key = childs[0]
                    data_array = getattr(datas, key)
                    childs.remove(key)
                    tmp_val = []
                    value_field = {}
                    numerical_order = 0

                    field_if = field.split('.IF.')

                    for data in data_array:
                        data_check = data
                        if len(field_if) > 1:
                            condition = field_if[1].split('=')
                            if len(condition):
                                childs_check = condition[0].split('.')
                                value = condition[1]
                            else:
                                childs_check = field_if[1].split('.')
                                value = True

                            for child in childs_check:
                                data_check = getattr(data_check, child)
                            if isinstance(data_check, bool):
                                value = bool(value)

                            if data_check != value:
                                continue

                        for child in childs:
                            if child == 'IF':
                                break
                            if child == 'numerical_order':
                                data = numerical_order + 1
                                numerical_order = data
                            elif child == "float_time":
                                hour, minute = divmod(data * 60, 60)
                                x_tmp = "%02d:%02d" % (hour, minute)
                                data = x_tmp
                            else:
                                image_value = data.id
                                data = getattr(data, child)
                            if isinstance(data, bytes):
                                image_value = str(image_value) + "." + field
                                break

                        if isinstance(data, (float, int)) == False and data == False:
                            data = ''
                        elif isinstance(data, bool):
                            data = ''
                        elif isinstance(data, datetime):
                            data = self._convert_datetime_usertz_to_utctz(data)
                        elif isinstance(data, date):
                            data = data.strftime(
                                self.env['res.lang'].search([('code', '=', self.env.user.lang)], limit=1).date_format)
                        elif isinstance(data, bytes):
                            data = image_value
                        else:
                            if isinstance(data, int) or isinstance(data, float):
                                data = self.format_number(data)
                            else:
                                data = str(data)
                        tmp_val.append(data)

                    value_field[field] = tmp_val
                    if key in populating_tables:
                        populating_tables[key].append(value_field)
                    else:
                        tmp_value = []
                        tmp_value.append(value_field)
                        populating_tables[key] = tmp_value
                elif childs[1] == 'user_sign' and len(childs)==2:
                    key = childs[0]
                    data_array = getattr(datas, key)
                    childs.remove(key)
                    tmp_val = []
                    value_field = {}
                    numerical_order = 0

                    field_if = field.split('.IF.')

                    for data in data_array:
                        data_check = data
                        if len(field_if) > 1:
                            condition = field_if[1].split('=')
                            if len(condition):
                                childs_check = condition[0].split('.')
                                value = condition[1]
                            else:
                                childs_check = field_if[1].split('.')
                                value = True

                            for child in childs_check:
                                data_check = getattr(data_check, child)
                            if isinstance(data_check, bool):
                                value = bool(value)

                            if data_check != value:
                                continue

                        for child in childs:
                            if child == 'IF':
                                break
                            if child == 'numerical_order':
                                data = numerical_order + 1
                                numerical_order = data
                            elif child == "float_time":
                                hour, minute = divmod(data * 60, 60)
                                x_tmp = "%02d:%02d" % (hour, minute)
                                data = x_tmp
                            else:
                                image_value = data.id
                                # digital_signature
                                data = getattr(data, 'digital_signature')
                            if isinstance(data, bytes):
                                image_value = str(image_value) + "." + field
                                break

                        if isinstance(data, (float, int)) == False and data == False:
                            data = ''
                        elif isinstance(data, bool):
                            data = ''
                        elif isinstance(data, datetime):
                            data = self._convert_datetime_usertz_to_utctz(data)
                        elif isinstance(data, date):
                            data = data.strftime(
                                self.env['res.lang'].search([('code', '=', self.env.user.lang)], limit=1).date_format)
                        elif isinstance(data, bytes):
                            data = image_value
                        else:
                            if isinstance(data, int) or isinstance(data, float):
                                data = self.format_number(data)
                            else:
                                data = str(data)
#                        print ('data ',data)
                        tmp_val.append(data)

                    value_field[field] = tmp_val
                    if key in populating_tables:
                        populating_tables[key].append(value_field)
                    else:
                        tmp_value = []
                        tmp_value.append(value_field)
                        populating_tables[key] = tmp_value
                elif childs[0] == 'line':
                    childs.remove(childs[0])
                    key = childs[0]
                    data_array = getattr(datas, key)
                    childs.remove(key)
                    tmp_val = []
                    value_field = {}
                    numerical_order = 0

                    field_if = field.split('.IF.')

                    for data in data_array:
                        data_check = data
                        if len(field_if) > 1:
                            condition = field_if[1].split('=')
                            if len(condition):
                                childs_check = condition[0].split('.')
                                value = condition[1]
                            else:
                                childs_check = field_if[1].split('.')
                                value = True

                            for child in childs_check:
                                data_check = getattr(data_check, child)
                            if isinstance(data_check, bool):
                                value = bool(value)

                            if data_check != value:
                                continue

                        for child in childs:
                            if child == 'IF':
                                break
                            if child == 'numerical_order':
                                data = numerical_order + 1
                                numerical_order = data
                            elif child == "float_time":
                                hour, minute = divmod(data * 60, 60)
                                x_tmp = "%02d:%02d" % (hour, minute)
                                data = x_tmp
                            else:
                                image_value = data.id
                                data = getattr(data, child)
                            if isinstance(data, bytes):
                                image_value = str(image_value) + "." + field
                                break

                        if isinstance(data, (float, int)) == False and data == False:
                            data = ''
                        elif isinstance(data, bool):
                            data = ''
                        elif isinstance(data, datetime):
                            data = self._convert_datetime_usertz_to_utctz(data)
                        elif isinstance(data, date):
                            data = data.strftime(
                                self.env['res.lang'].search([('code', '=', self.env.user.lang)], limit=1).date_format)
                        elif isinstance(data, bytes):
                            data = image_value
                        else:
                            if isinstance(data, int) or isinstance(data, float):
                                data = self.format_number(data)
                            else:
                                data = str(data)
                        tmp_val.append(data)

                    value_field[field] = tmp_val
                    if key in populating_tables:
                        populating_tables[key].append(value_field)
                    else:
                        tmp_value = []
                        tmp_value.append(value_field)
                        populating_tables[key] = tmp_value


                elif childs[0] == 'line_fix':
                    childs.remove(childs[0])
                    key = childs[0]
                    data_array = getattr(datas, key)
                    childs.remove(key)
                    key = childs[0]
                    for data in data_array:
                        data_key = getattr(data, key)
                        if data_key.id == int(childs[1]):
                            if childs[2] != 'label':
                                value = getattr(data, childs[2])
                                if data._name == 'tests' and data.parameter_characteristic_id.unit_id.id != False:
                                    unit = data.parameter_characteristic_id.unit_id.name
                                    value = value + ' ' + unit
                                if isinstance(data, datetime):
                                    value = self._convert_datetime_usertz_to_utctz(value)
                                elif isinstance(value, date):
                                    value = value.strftime(self.env['res.lang'].search([('code', '=', self.env.user.lang)],
                                                                                       limit=1).date_format)
                                elif isinstance(value, bool):
                                    if value == False:
                                        value = ''
                                    else:
                                        value = str(value)
                                elif isinstance(value, str):
                                    value = str(value)

                                elif isinstance(value, bytes):
                                    value = str(data.id) + "." + field
                                elif isinstance(value, object):
                                    if len(value) > 1:
                                        value_str = ''
                                        length = 1
                                        for val in value:
                                            if len(value) == length:
                                                value_str += getattr(val, 'name')
                                            else:
                                                value_str += getattr(val, 'name') + ', '
                                                length += 1
                                        value = value_str
                                    else:
                                        value = getattr(value, 'name')
                                else:
                                    if isinstance(value, int) or isinstance(value, float):
                                        value = self.format_number(value)
                                    else:
                                        value = str(value)
                                if field in simple_merge:
                                    if len(childs) == 4 and childs[3] == 'merge':
                                        pass
                                    else:
                                        simple_merge[field] = simple_merge[field] + '           ' + value
                                else:
                                    simple_merge[field] = value
                            else:
                                if data._name == 'tests':
                                    values = data.parameter_id.get_field_translations('name')
                                    for v in values[0]:
                                        if len(childs) < 4:
                                            if v['lang'] == self.env.lang:
                                                value = v['value']
                                        else:
                                            if childs[3] == v['lang']:
                                                value = v['value']
                                    simple_merge[field] = value

                else:
                    if len(childs) <= 0:
                        continue
                    tmp_logic = childs[len(childs)-1]
                    if tmp_logic == 'sum':
                        data_array = getattr(datas, childs[0])
                        sum = 0
                        if len(childs) == 3:
                            for data in data_array:
                                value = getattr(data, childs[1])
                                sum += value
                            simple_merge[field] = self.format_number(sum)
                        elif len(childs) == 4:
                            data_array = getattr(data_array, childs[1])
                            for data in data_array:
                                value = getattr(data, childs[2])
                                sum += value
                            simple_merge[field] = self.format_number(sum)
                        else:
                            simple_merge[field] = 0
                    elif tmp_logic == 'count':
                        data_array = getattr(datas, childs[0])
                        count = len(data_array)
                        simple_merge[field] = self.format_number(count)
                    elif tmp_logic == 'sum_number2word':
                        data_array = getattr(datas, childs[0])
                        sum = 0

                        for data in data_array:
                            value = getattr(data, childs[1])
                            sum += value
                        num_to_char = self.num2word(sum)
                        simple_merge[field] = num_to_char
                    elif tmp_logic == 'width':
                        simple_merge[field] = field
                    else:
                        data = datas
                        for child in childs:
                            data = getattr(data,child)
                        simple_merge[field] = str(data)
                        try:
                            simple_merge[field] = self.format_number(data)
                        except Exception:
                            simple_merge[field] = str(data)

        document.merge(**simple_merge)
        for key in populating_tables:
            value = populating_tables[key]
            list = []
            anchor = ''
            number = 0
            if number == 0:
                for k in value[0]:
                    val = value[0][k]
                    number = len(val)
                    break
            for i in range(number):
                dict = {}
                for val in value:
                    for k in val:
                        v = val[k]
                        dict[k] = v[i]
                        if anchor == '':
                            anchor = k
                        break
                list.append(dict)
            document.merge_rows(anchor, list)
        for field in document.get_merge_fields():
            document.merge(**{field: ''})
        mem_zip = BytesIO()
        with ZipFile(mem_zip, 'w', ZIP_DEFLATED) as output:
            for zi in document.zip.filelist:
                if zi in document.parts:
                    xml = etree.tostring(document.parts[zi].getroot())
                    output.writestr(zi.filename, xml)
                elif zi == document._settings_info:
                    xml = etree.tostring(document.settings.getroot())
                    output.writestr(zi.filename, xml)
                else:
                    output.writestr(zi.filename, document.zip.read(zi))

        tempfile_docx = tempfile.NamedTemporaryFile(suffix='docx')
        with open(tempfile_docx.name, "wb+") as f:
            f.write(mem_zip.getbuffer())
        f.close()
        doc = Document(tempfile_docx.name)
        for block in self._iter_block_items(doc):
            if isinstance(block, Table):
                self._replace_table_cell_with_image(block, datas, fields)
            else:
                self._match_and_replace(block, datas, fields)
        doc.save(tempfile_docx.name)

        with open(tempfile_docx.name, 'rb') as docx_file:
            docx_binary_value = docx_file.read()
        if self.type_export == 'docx':
            return docx_binary_value
        else:
            return self.convert_docx_to_pdf(docx_binary_value)

    def _get_suffix(self):
        return str(self.file_template_name).split(".")[-1]

    def render_doc_doc(self, res_ids=None, data=None):
        suffix = self._get_suffix()
        docx = self.export_doc_by_template(datas=res_ids[0], file_name_export=self.print_report_name,suffix=suffix)
        if self.type_export == 'pdf':
            suffix = 'pdf'
        return docx, suffix

    def _convert_datetime_usertz_to_utctz(self, value):
        # convert back from user's timezone to UTC
        tz_name = self.env.context.get('tz') or self.env.user.tz
        if tz_name:
            try:
                user_tz = pytz.timezone(tz_name)
                dt = value.astimezone(user_tz)
            except Exception:
                _logger.warn(
                    "Failed to convert the value for a field of the model"
                    " %s back from the user's timezone (%s) to UTC",
                    'ir.actions.report', tz_name,
                    exc_info=True)
        return dt.strftime(self.env['res.lang'].search([('code','=', self.env.user.lang)], limit=1).date_format +" " + self.env['res.lang'].search([('code','=', self.env.user.lang)], limit=1).time_format)

    def format_number(self, value):
        value=round(value,2)
        lang = self.env['res.lang'].search([('code', '=', self.env.user.lang)])
        value = "{:,}".format(value)
        value = value.replace('.', 'x')
        value = value.replace(',', lang.thousands_sep)
        value = value.replace('x', lang.decimal_point)
        return value

    def num2word(self, amount):
        self.ensure_one()

        def _num2words(number, lang):
            try:
                return num2words(number, lang=lang).title()
            except NotImplementedError:
                return num2words(number, lang='en').title()

        if num2words is None:
            logging.getLogger(__name__).warning("The library 'num2words' is missing, cannot render textual amounts.")
            return ""

        formatted = "%.{0}f".format(0) % amount
        parts = formatted.partition('.')
        integer_value = abs(int(parts[0]))

        lang_code = self.env.context.get('lang') or self.env.user.lang
        amount_words = tools.ustr('{amt_value} {amt_word}').format(
            amt_value=_num2words(integer_value, lang=lang_code),
            amt_word=''
        )
        return amount_words

    def convert_docx_to_pdf(self, source, timeout=15):
        filetmp = tempfile.NamedTemporaryFile(suffix='.docx')
        filetmp.write(source)
        filetmp.seek(0)
        dir = tempfile.gettempdir()
        args = [self.libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', dir, filetmp.name]
        process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
        filename = re.search('-> (.*?) using filter', process.stdout.decode())
        filename = filename.group(1)
        with open(filename, 'rb') as pdf_document:
            pdf_content = pdf_document.read()
        return pdf_content

    def libreoffice_exec(self):
        # TODO: Provide support for more platforms
        if sys.platform == 'darwin':
            return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
        return 'libreoffice'

    def _render_qweb_pdf(self, report_ref, res_ids=None, data=None):
        if self.populating_ms_word_template != True:
            result = super(IrActionsReport, self)._render_qweb_pdf(report_ref, res_ids, data)
        else:
            datas = self.env[self.model].browse(res_ids[0])
            result = self.render_doc_doc(datas,data)
        return result

    def _iter_block_items(self, parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _match_and_replace(self, block, record, fields):
        # print ('block.text ',block.text)
        if 'res.users' in   block.text: # гарын үсэг
            # print ('block.text ',block.text)
            users='users'
            user_id=block.text.split('res.users(')[1]
            user_id=user_id.split(',)')[0]
            # exec(f"{users} = {block.text}")
            users=self.env['res.users'].browse(int(user_id))
            image = users.digital_signature #getattr(record, field[0], b'')
            if not image and users.digital_signature_from_file:
                image = users.digital_signature_from_file 
            if isinstance(block, _Cell):
                block = block.paragraphs[0]
            if isinstance(image, bytes):
                image = base64_to_image(image)
            else:
                if isinstance(image, str):
                    image = self.barcode('Code128', image)
                    image = base64_to_image(base64.b64encode(image).decode('utf-8'))
                # else:
                #     aa=getattr(image, field[1], b'')
                #     # print ('aa ',aa)
                #     image = base64_to_image(aa)
            if image:
                tempfile_png = tempfile.NamedTemporaryFile(suffix='.png')
                image.save(tempfile_png.name)
                block.text = ''
                block.add_run().add_picture(tempfile_png.name, width=Inches(2))  
        elif '.user_sign' in   block.text: # гарын үсэг
            matched_field = list(filter(lambda field: field in block.text.strip(), fields))
            matched_field=[]
            for f in fields:
                ff='.'+f #+'.user_sign'
                # print ('ff ',ff)
                if ff in block.text.strip():
                    matched_field.append(f)
            if matched_field and not hasattr(record, matched_field[0]):
                field = matched_field[0].split('.')
            else:
                field = matched_field
            # if 'user_sign' in  block.text and matched_field: # гарын үсэг  
            if len(matched_field)>1:
                for matched in  matched_field: # origin, origin_user_id gm davhardsan ug orson bol
                    if 'user_sign' in matched:
                        matched_field = matched
            # print ('matched_field ',matched_field)
            if matched_field and hasattr(record, field[0]):
                image = getattr(record, field[0], b'')
                _logger.info(u'record-------------3%s !'%(record))
                _logger.info(u'field[0]-------------3%s !'%(field[0]))
                _logger.info(u'image-------------3%s !'%(image))
                if isinstance(image, bytes):
                    image = base64_to_image(image)
                else:
                    if isinstance(image, str):
                        # image = self.barcode('Code128', image)
                        # image = base64_to_image(base64.b64encode(image).decode('utf-8'))
                        try:
                            image = base64_to_image(base64.b64encode(image).decode('utf-8'))
                        except Exception:
                            image = base64_to_image(base64.b64encode(image.encode('utf-8')).decode('utf-8'))
                    else:
                        fff=field[1]
                        if field[1]=='user_sign':
                            fff='digital_signature'
                        image = base64_to_image(getattr(image, fff, b''))
                tempfile_png = tempfile.NamedTemporaryFile(suffix='.png')
                image.save(tempfile_png.name)
                block.text = ''
                if isinstance(block, _Cell):
                    block = block.paragraphs[0]
                if len(field) == 3:
                    # block.add_run().add_picture(tempfile_png.name, width=Inches(int(field[1])))
                    block.add_run().add_picture(tempfile_png.name, width=Inches(2))
                elif len(field) == 2:
                    block.add_run().add_picture(tempfile_png.name, width=Inches(2))            
            return 
        matched_field = list(filter(lambda field: field in block.text.strip(), fields))
        if matched_field and not hasattr(record, matched_field[0]):
            field = matched_field[0].split('.')
        else:
            field = matched_field

        # if 'user_sign' in  block.text and matched_field: # гарын үсэг  
        if matched_field and hasattr(record, field[0]):
            image = getattr(record, field[0], b'')
            if isinstance(image, bytes):
                image = base64_to_image(image)
            else:
                if isinstance(image, str):
                    image = self.barcode('Code128', image)
                    image = base64_to_image(base64.b64encode(image).decode('utf-8'))
                else:
                    fff=field[1]
                    if field[1]=='user_sign':
                        fff='digital_signature'
                    # print ('fff ',fff)
                    image = base64_to_image(getattr(image, fff, b''))
            tempfile_png = tempfile.NamedTemporaryFile(suffix='.png')
            image.save(tempfile_png.name)
            block.text = ''
            if isinstance(block, _Cell):
                block = block.paragraphs[0]
            if len(field) == 3:
                block.add_run().add_picture(tempfile_png.name, width=Inches(int(field[1])))
                # block.add_run().add_picture(tempfile_png.name, width=Inches(2))
            elif len(field) == 2:
                block.add_run().add_picture(tempfile_png.name, width=Inches(2))
            elif len(field) == 4:
                if field[2] == 'cm':
                    width = int(field[1]) / 2.54
                elif field[2] == 'mm':
                    width = int(field[1]) / 25.4
                elif field[2] == 'px':
                    width = int(field[1]) / 96
                else:
                    width = int(field[2])
                block.add_run().add_picture(tempfile_png.name, width=Inches(width))
            elif len(field) == 5:
                if field[3] == 'cm':
                    width = int(field[2]) / 2.54
                elif field[3] == 'mm':
                    width = int(field[2]) / 25.4
                elif field[3] == 'px':
                    width = int(field[2]) / 96
                else:
                    width = int(field[2])
                block.add_run().add_picture(tempfile_png.name, width=Inches(width))
            else :
                block.add_run().add_picture(tempfile_png.name)
        else:
            if matched_field:
                childs = block.text.split('.')
                block.text = ''
                data_ids = getattr(record, childs[2], b'')
                for data in data_ids:
                    if data.id == int(childs[0]):
                        if childs[1] == 'line_fix':
                            image = base64_to_image(getattr(data, childs[5], b''))
                            tempfile_png = tempfile.NamedTemporaryFile(suffix='.png')
                            image.save(tempfile_png.name)
                            if isinstance(block, _Cell):
                                block = block.paragraphs[0]
                            if childs[7] == 'width':
                                block.add_run().add_picture(tempfile_png.name, width=Inches(int(childs[4])))
                            elif childs[8] == 'width':
                                if childs[7] == 'cm':
                                    width = int(field[5]) / 2.54
                                elif childs[7] == 'mm':
                                    width = int(field[5]) / 25.4
                                elif childs[7] == 'px':
                                    width = int(field[5]) / 96
                                else:
                                    width = int(field[5])
                                block.add_run().add_picture(tempfile_png.name, width=Inches(width))
                            else:
                                block.add_run().add_picture(tempfile_png.name)
                            break
                        else:
                            image = base64_to_image(getattr(data, childs[3], b''))
                            tempfile_png = tempfile.NamedTemporaryFile(suffix='.png')
                            image.save(tempfile_png.name)
                            if isinstance(block, _Cell):
                                block = block.paragraphs[0]
                            if childs[5] == 'width':
                                block.add_run().add_picture(tempfile_png.name, width=Inches(int(1)))
                            elif childs[6] == 'width':
                                if childs[5] == 'cm':
                                    width = int(field[1]) / 2.54
                                elif childs[5] == 'mm':
                                    width = int(field[1]) / 25.4
                                elif childs[5] == 'px':
                                    width = int(field[1]) / 96
                                else:
                                    width = int(field[1])
                                block.add_run().add_picture(tempfile_png.name, width=Inches(width))
                            else:
                                block.add_run().add_picture(tempfile_png.name)
                            break

    def _replace_table_cell_with_image(self, table, record, fields):
        for (i, row) in enumerate(table.rows):
            for (j, cell) in enumerate(row.cells):
                # print ('cell ',cell.text)
                self._match_and_replace(cell, record, fields)